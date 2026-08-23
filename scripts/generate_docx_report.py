import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DOCX_PATH = r"D:\Magang BPS\Project\bps_ai_mobile\LAPORAN_BPS_AI_ASSISTANT_MOBILE.docx"

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    """Set clean formal borders for table."""
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

def build_docx():
    doc = docx.Document()
    
    # ── Page Setup (A4, 2.54cm margins) ──
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Badan Pusat Statistik Provinsi Sulawesi Tengah — Laporan Proyek BPS AI Assistant Mobile"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.name = "Times New Roman"
        hp.style.font.size = Pt(8.5)
        hp.style.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Dokumentasi Teknis & Laporan Hasil Pengembangan Sistem"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.style.font.name = "Times New Roman"
        fp.style.font.size = Pt(8.5)
        fp.style.font.color.rgb = RGBColor(128, 128, 128)

    # ── Styles Setup ──
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.italic = True
        run.font.color.rgb = RGBColor(50, 50, 50)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_body(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(5)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        p.add_run(text)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        p.add_run(text)
        return p

    # ── JUDUL DOKUMEN ──
    add_title("LAPORAN PROYEK PENGEMBANGAN SISTEM")
    add_title("BPS AI ASSISTANT MOBILE")
    add_subtitle("Inovasi Asisten Statistik Cerdas Berbasis Generative AI dan Integrasi BPS Web API")

    # ── TABEL IDENTITAS ──
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    meta_data = [
        ("Pengembang", "Tim Asisten Pranata Komputer"),
        ("Unit Kerja / Instansi", "Badan Pusat Statistik Provinsi Sulawesi Tengah"),
        ("Status Proyek", "Fase 1 Selesai (Live Production Ready)"),
        ("Akses Layanan Cloud", "https://bps-ai-backend.vercel.app / Installer: app-release.apk")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        r1 = p_lbl.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10)
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(2)
        r2 = p_val.add_run(val)
        r2.font.size = Pt(10)
        
        set_cell_background(cell_lbl, "F2F2F2")
        set_cell_background(cell_val, "FAFAFA")
        
    set_table_borders(meta_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── BAB I: RINGKASAN EKSEKUTIF ──
    add_h1("1. RINGKASAN EKSEKUTIF")
    add_body(
        "Aplikasi BPS AI Assistant Mobile adalah sistem diseminasi data statistik berbasis kecerdasan buatan "
        "(Artificial Intelligence) yang dirancang untuk mempermudah pimpinan instansi, akademisi, peneliti, pemerintah daerah, "
        "dan masyarakat luas dalam mengakses serta memahami data resmi Badan Pusat Statistik (BPS) secara cepat, interaktif, dan akurat."
    )
    add_body(
        "Sistem ini menerapkan metode Retrieval-Augmented Generation (RAG) dua tahap yang terhubung langsung secara real-time "
        "ke portal resmi BPS Web API (webapi.bps.go.id). Dengan integrasi ini, asisten AI tidak menggunakan data karangan "
        "atau data usang, melainkan menyajikan data aktual yang bersumber langsung dari Berita Resmi Statistik (BRS) dan publikasi BPS."
    )
    add_body(
        "Pada Fase 1, pengembangan telah diselesaikan secara menyeluruh yang mencakup: (1) Gateway backend FastAPI yang beroperasi aktif "
        "di server cloud Vercel, (2) Mekanisme ketahanan multi-model (Gemini 3.7, 3.6, dan 3.5 Flash) untuk mencegah gangguan server 503, "
        "serta (3) Berkas installer aplikasi Android Release APK (app-release.apk) yang telah dilengkapi izin akses internet penuh."
    )

    # ── BAB II: LATAR BELAKANG & TUJUAN STRATEGIS ──
    add_h1("2. LATAR BELAKANG DAN TUJUAN STRATEGIS")
    add_h2("2.1 Latar Belakang Masalah")
    add_bullet("Publik dan pemangku kepentingan sering menghadapi kesulitan dalam mencari indikator spesifik pada portal web yang memuat ribuan tabel data.", "1. Kompleksitas Penelusuran Data: ")
    add_bullet("Pengguna kerap membutuhkan penjelasan naratif ringkas mengenai makna indikator ekonomi dan sosial, bukan sekadar deretan angka mentah.", "2. Kebutuhan Konteks & Narasi: ")
    add_bullet("Diperlukan sarana diseminasi resmi yang menjamin bahwa seluruh data yang beredar di masyarakat bersumber secara sahih dari BPS RI.", "3. Mitigasi Misinformasi: ")

    add_h2("2.2 Tujuan Pengembangan")
    add_bullet("Menyediakan sarana konsultasi cerdas mengenai definisi dan metodologi indikator statistik (Inflasi, PDRB, Kemiskinan, IPM, TPT, NTP).", "• Literasi Statistik: ")
    add_bullet("Menyajikan data rilis terkini secara instan lengkap dengan tautan unduh dokumen resmi (PDF Berita Resmi Statistik).", "• Akselerasi Diseminasi: ")
    add_bullet("Membantu pimpinan BPS dan instansi pemerintah dalam mengakses data indikator makro secara cepat untuk perumusan kebijakan berbasis bukti (data-driven policy).", "• Dukungan Kebijakan: ")

    # ── BAB III: ARSITEKTUR SISTEM & SPESIFIKASI TEKNOLOGI ──
    add_h1("3. ARSITEKTUR SISTEM DAN SPESIFIKASI TEKNOLOGI")
    add_body(
        "Sistem dibangun dengan memisahkan antarmuka pengguna (Frontend Mobile), lapisan pemrosesan integrasi (Backend Gateway), "
        "model bahasa besar (Large Language Model), serta basis data resmi BPS Web API."
    )

    # Tabel Tech Stack
    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_table.autofit = False
    
    headers = ["Komponen", "Teknologi", "Fungsi dan Peran"]
    for i, h in enumerate(headers):
        cell = tech_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        set_cell_background(cell, "E0E0E0")
        
    tech_rows = [
        ("Frontend Mobile", "Flutter SDK 3.38.7 / Dart 3.10.7", "Antarmuka aplikasi mobile Android, tema standar BPS, dan kartu UI terstruktur."),
        ("Backend Gateway", "FastAPI (Python 3.13) + Uvicorn", "Logika RAG, ekstraksi kata kunci, sanitasi data, dan manajemen endpoint API."),
        ("AI Engine", "Google Gemini 3.7 / 3.6 / 3.5 Flash", "Pemrosesan bahasa alami Indonesia, ekstraksi entitas, dan penyusunan narasi data."),
        ("Sumber Data Resmi", "BPS Web API (webapi.bps.go.id)", "Penyedia data statistik resmi (BRS, tabel statis, subjek statistik, publikasi)."),
        ("Infrastruktur Cloud", "Vercel Serverless Edge Cloud", "Server backend aktif online 24/7 dengan protokol keamanan HTTPS global."),
        ("Paket Distribusi", "Android Release APK (45.3 MB)", "Installer aplikasi mandiri (standalone) dengan izin akses jaringan penuh.")
    ]
    
    col_widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    for row_idx, data in enumerate(tech_rows, start=1):
        row = tech_table.rows[row_idx]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.bold = True
            if row_idx % 2 == 0:
                set_cell_background(cell, "F9F9F9")
                
    set_table_borders(tech_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_h2("3.1 Alur Kerja Penarikan Data (Retrieval-Augmented Generation)")
    add_bullet("Pengguna memasukkan pertanyaan melalui aplikasi Android (misal: 'Berapa inflasi Indonesia pada Juli 2026?').", "1. Input Pertanyaan: ")
    add_bullet("Gateway backend memproses pertanyaan ke model AI untuk mengekstrak kata kunci ('inflasi'), model data ('pressrelease'), dan kode wilayah ('0000' untuk nasional).", "2. Ekstraksi Intent & Entitas: ")
    add_bullet("Backend mengirim permintaan data secara asinkron ke server BPS Web API menggunakan API Key resmi.", "3. Query Data Riil BPS: ")
    add_bullet("Server BPS mengembalikan metadata angka rilis (2,88%) dan berkas unduhan PDF resmi.", "4. Respons Data BPS: ")
    add_bullet("AI menyusun ringkasan naratif formal berdasarkan data yang diterima, memetakan data payload angka, dan melampirkan judul rilis resmi.", "5. Sintesis Narasi: ")
    add_bullet("Aplikasi Flutter menampilkan kartu statistik angka terstruktur beserta tombol langsung untuk membuka berkas PDF rilis resmi BPS.", "6. Penyajian Antarmuka: ")

    # ── BAB IV: FITUR YANG TELAH SELESAI ──
    add_h1("4. FITUR YANG TELAH SELESAI DIIMPLEMENTASIKAN")
    add_bullet("Mampu menjawab pertanyaan seputar inflasi, kemiskinan, ketenagakerjaan, pertanian (NTP), dan kependudukan berbasis data rilis resmi BPS.", "1. Integrasi Data Riil BPS (Live RAG): ")
    add_bullet("Menampilkan kartu visual ringkas yang memuat nama indikator, nilai capaian, cakupan wilayah, periode data, dan tombol tautan dokumen PDF rilis resmi.", "2. Kartu Statistik Terstruktur & Sitasi PDF: ")
    add_bullet("Mekanisme pengalihan otomatis berantai (Gemini 3.7 -> 3.5 -> 3.6 -> Flash Latest) untuk memastikan ketersediaan sistem tanpa kendala server overloaded saat pengujian.", "3. Sistem Ketahanan Multi-Model (Anti-503): ")
    add_bullet("Penerapan palet warna standar institusi BPS (Biru, Oranye, Hijau) dengan optimasi tata letak yang bebas dari peringatan render overflow pada seluruh resolusi layar.", "4. Desain Antarmuka Standar BPS: ")
    add_bullet("Backend aktif 24 jam di cloud Vercel dan aplikasi telah dikompilasi menjadi installer mandiri yang siap dipasang pada perangkat pimpinan/pengguna.", "5. Kesiapan Deployment Cloud & File APK: ")

    # ── BAB V: HASIL PENGUJIAN SISTEM ──
    add_h1("5. HASIL PENGUJIAN SISTEM DAN VALIDASI DATA")
    add_body("Berikut adalah dokumentasi hasil pengujian aktual pada beberapa skenario pertanyaan:")

    test_table = doc.add_table(rows=4, cols=3)
    test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_table.autofit = False
    
    t_headers = ["Skenario Pertanyaan", "Proses Pengolahan Backend", "Hasil Respon Sistem"]
    for i, h in enumerate(t_headers):
        cell = test_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        set_cell_background(cell, "E0E0E0")
        
    test_cases = [
        ("Data Angka Indikator:\n'Berapa inflasi Indonesia pada Juli 2026?'",
         "Query BPS Web API model pressrelease dengan keyword 'inflasi' dan domain '0000'.",
         "Status: 200 OK\nIndikator: Inflasi y-on-y\nNilai: 2,88 persen (IHK: 111,73)\nPeriode: Juli 2026\nTersedia tombol unduh PDF resmi BRS."),
        
        ("Data Proyeksi Kependudukan:\n'Berapa jumlah penduduk tahun 2025?'",
         "Identifikasi topik demografi dan perujukan ke data resmi Proyeksi Sensus Penduduk SP2020.",
         "Status: 200 OK\nNilai Proyeksi: 284,43 Juta Jiwa\nRujukan: Proyeksi Penduduk 2020-2050 BPS-Bappenas (Skenario Tren)."),
        
        ("Definisi & Konsep:\n'Apa itu inflasi?'",
         "Identifikasi pertanyaan pengetahuan statistik dan penyusunan narasi edukasi konsep IHK.",
         "Status: 200 OK\nNarasi: Penjelasan konsep kenaikan harga umum yang diukur melalui Indeks Harga Konsumen di kota-kota pantauan BPS.")
    ]
    
    t_widths = [Inches(2.0), Inches(2.2), Inches(2.3)]
    for row_idx, data in enumerate(test_cases, start=1):
        row = test_table.rows[row_idx]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = t_widths[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if col_idx == 0:
                r.bold = True
            if row_idx % 2 == 0:
                set_cell_background(cell, "F9F9F9")
                
    set_table_borders(test_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── BAB VI: RENCANA PENGEMBANGAN LANJUTAN ──
    add_h1("6. RENCANA DAN ROADMAP PENGEMBANGAN LANJUTAN")
    add_body(
        "Untuk mentransformasikan aplikasi BPS AI Assistant Mobile menjadi produk unggulan diseminasi data statistik nasional, "
        "Tim Asisten Pranata Komputer BPS Provinsi Sulawesi Tengah merekomendasikan tahapan pengembangan fitur lanjutan sebagai berikut:"
    )

    roadmap_table = doc.add_table(rows=7, cols=3)
    roadmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roadmap_table.autofit = False
    
    r_headers = ["Rencana Fitur", "Deskripsi Pengembangan Teknis", "Manfaat bagi Pimpinan / BPS"]
    for i, h in enumerate(r_headers):
        cell = roadmap_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        set_cell_background(cell, "E0E0E0")
        
    roadmap_rows = [
        ("Visualisasi Grafik Interaktif", "Integrasi modul grafik garis (Line Chart) dan batang (Bar Chart) interaktif otomatis untuk data deret waktu (time-series).", "Pimpinan dapat memahami arah tren indikator ekonomi/sosial secara visual dalam hitungan detik."),
        ("Komparasi Antar-Wilayah", "Tabel perbandingan data berdampingan (side-by-side comparison) antar-provinsi atau kabupaten/kota beserta perhitungan selisih angka.", "Mempermudah evaluasi disparitas pembangunan antar-daerah bagi pemerintah daerah."),
        ("Smart Policy Bridging", "Pengaitan pertanyaan program pemerintah (MBG, IKN, Bansos) ke dataset BPS pendukung (Susenas, Prevalensi PoU, Kemiskinan).", "Menonjolkan peran strategis BPS sebagai rujukan data utama pengawalan kebijakan nasional."),
        ("Voice Assistant Multimodal", "Penambahan fitur input suara (Speech-to-Text) dan pembaca audio ringkasan statistik (Text-to-Speech) dalam Bahasa Indonesia.", "Meningkatkan aksesibilitas bagi pimpinan saat mobile serta pengguna disabilitas netra."),
        ("Ekspor PDF & Berbagi ke WhatsApp", "Fitur ekspor lembar Ringkasan Eksekutif resmi format PDF satu klik dan format kartu ringkasan untuk dibagikan ke WhatsApp.", "Mempercepat diseminasi data untuk bahan paparan rapat pimpinan dan siaran pers."),
        ("Live BRS Feed & Dashboard Indikator", "Halaman dashboard pemantauan indikator makro strategis dan feed rilis Berita Resmi Statistik harian.", "Menjadi portal pemantauan indikator strategis harian pimpinan BPS.")
    ]
    
    r_widths = [Inches(1.8), Inches(2.5), Inches(2.2)]
    for row_idx, data in enumerate(roadmap_rows, start=1):
        row = roadmap_table.rows[row_idx]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = r_widths[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if col_idx == 0:
                r.bold = True
            if row_idx % 2 == 0:
                set_cell_background(cell, "F9F9F9")
                
    set_table_borders(roadmap_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── BAB VII: KESIMPULAN ──
    add_h1("7. KESIMPULAN DAN PENUTUP")
    add_body(
        "Pengembangan BPS AI Assistant Mobile (Fase 1) oleh Tim Asisten Pranata Komputer BPS Provinsi Sulawesi Tengah "
        "telah berhasil diselesaikan dengan baik, berjalan stabil pada infrastruktur server cloud, dan memenuhi seluruh standar "
        "keandalan integrasi data resmi BPS RI."
    )
    add_body(
        "Aplikasi ini siap dipresentasikan di hadapan pimpinan sebagai bukti nyata inovasi digitalisasi dan modernisasi pelayanan "
        "data statistik Badan Pusat Statistik menuju era keterbukaan data yang inklusif, modern, dan akurat."
    )
    
    doc.save(DOCX_PATH)
    print("Formal Word Document (.docx) generated at:", DOCX_PATH)

if __name__ == "__main__":
    build_docx()
